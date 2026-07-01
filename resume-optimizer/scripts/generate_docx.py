#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历Word文档生成脚本
功能：将MD简历转换为Word文档
用法：python generate_docx.py <input.md> <output.docx>

依赖：python-docx 库
"""

import sys
import os
import re


def add_run_spacing_to_rpr(run, line_spacing_mult):
    """在run的rPr元素中添加spacing"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # 获取或创建rPr元素
    rPr = run._r.get_or_add_rPr()
    
    # 创建spacing元素
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(line_spacing_mult * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    
    # 移除旧的spacing元素
    for child in rPr:
        if child.tag == qn('w:spacing'):
            rPr.remove(child)
    
    # 在rPr的末尾添加spacing
    rPr.append(spacing)


def md_to_docx(md_content, output_file):
    """
    将Markdown内容转换为Word文档
    - 非加粗字体：单倍行距
    - 加粗字体：1.5倍行距
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("错误：未安装 python-docx 库，请运行: pip install python-docx", file=sys.stderr)
        return False
    
    doc = Document()
    
    # 设置页面格式为A4
    section = doc.sections[0]
    section.page_width = Pt(595.276)
    section.page_height = Pt(841.89)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(8.5)
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    lines = md_content.split('\n')
    divider_added = False
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            continue
        
        # 处理一级标题（姓名）- 加粗，1.5倍行距，10号字
        if stripped.startswith('# ') and not stripped.startswith('##'):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            add_run_spacing_to_rpr(run, 1.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        
        # 处理二级标题 - 加粗，1.5倍行距，专业蓝色
        elif stripped.startswith('## ') and not stripped.startswith('###'):
            title_text = stripped[3:]
            divider_added = False
            
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204)  # 专业蓝色 #0066CC
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            add_run_spacing_to_rpr(run, 1.5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
        
        # 处理三级标题 - 加粗，1.5倍行距
        elif stripped.startswith('### '):
            title_text = stripped[4:]
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            add_run_spacing_to_rpr(run, 1.5)  # 加粗使用1.5倍行距
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(0)
        
        # 处理分隔线 - 使用边框线
        elif stripped.startswith('---'):
            if not divider_added:
                p = doc.add_paragraph()
                run = p.add_run('')
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                add_run_spacing_to_rpr(run, 1.0)  # 单倍行距
                
                # 添加下边框
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                pBdr.append(bottom)
                pPr.append(pBdr)
                
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(0)
                divider_added = True
        
        # 处理列表项
        elif stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:]
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
            
            p = doc.add_paragraph()
            
            # 项目符号 - 非加粗，单倍行距
            run = p.add_run('● ')
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            add_run_spacing_to_rpr(run, 1.0)  # 单倍行距
            
            # 内容 - 非加粗，单倍行距
            run = p.add_run(content)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            add_run_spacing_to_rpr(run, 1.0)  # 单倍行距
            
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        
        # 处理普通段落
        else:
            divider_added = False
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            add_run_spacing_to_rpr(run, 1.0)  # 单倍行距
            
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.first_line_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    try:
        doc.save(output_file)
        print(f"Word文档已生成：{output_file}")
        return True
    except Exception as e:
        print(f"错误：保存Word文档时出错: {str(e)}", file=sys.stderr)
        return False


def main():
    if len(sys.argv) < 3:
        print("用法：python generate_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"错误：输入文件不存在 - {input_file}", file=sys.stderr)
        sys.exit(1)
    
    md_content = None
    encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312']
    
    for encoding in encodings:
        try:
            with open(input_file, 'r', encoding=encoding) as f:
                md_content = f.read()
            break
        except UnicodeDecodeError:
            continue
    
    if md_content is None:
        print(f"错误：无法读取文件编码 - {input_file}", file=sys.stderr)
        sys.exit(1)
    
    if not md_content.strip():
        print(f"错误：文件内容为空 - {input_file}", file=sys.stderr)
        sys.exit(1)
    
    success = md_to_docx(md_content, output_file)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
