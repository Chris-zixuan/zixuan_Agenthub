#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历Word文档读取脚本（使用 python-docx 库）
功能：读取 .docx 文件，提取文本内容和表格结构
用法：python read_docx.py <input_file.docx> [output_file.txt]
      python read_docx.py <input_file.docx> --json [output_file.json]
      python read_docx.py <input_file.docx> --structure [output_file.json]

依赖：python-docx 库
"""

import sys
import os
import re
import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def extract_text_from_paragraph(paragraph):
    """提取段落文本，处理特殊格式"""
    text_parts = []
    
    for run in paragraph.runs:
        text = run.text
        if text:
            # 保留制表符和换行符
            text = text.replace('\t', '\t').replace('\n', '\n')
            text_parts.append(text)
    
    return ''.join(text_parts)


def parse_docx(file_path):
    """
    读取Word文档内容，返回纯文本和结构化数据
    支持 .docx 格式（使用 python-docx 库）
    """
    if not os.path.exists(file_path):
        return {"success": False, "error": f"文件不存在: {file_path}", "content": ""}
    
    if not file_path.lower().endswith('.docx'):
        return {"success": False, "error": "仅支持 .docx 格式", "content": ""}
    
    try:
        doc = Document(file_path)
        
        paragraphs_text = []
        tables_content = []
        lists_content = []
        current_list = []
        in_list = False
        
        for paragraph in doc.paragraphs:
            text = extract_text_from_paragraph(paragraph).strip()
            
            if not text:
                continue
            
            # 获取段落样式
            style_name = paragraph.style.name if paragraph.style else ""
            is_list = False
            list_level = 0
            
            # 检查是否为列表项
            if paragraph._p.pPr is not None:
                num_pr = paragraph._p.pPr.numPr
                if num_pr is not None:
                    is_list = True
                    # 获取列表级别
                    ilvl = num_pr.xpath('.//w:ilvl')
                    if ilvl:
                        list_level = int(ilvl[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0'))
            
            if is_list:
                if not in_list:
                    in_list = True
                    current_list = []
                current_list.append((list_level, text))
            else:
                if in_list:
                    lists_content.append(current_list)
                    in_list = False
                    current_list = []
                
                # 根据样式判断标题级别
                if style_name:
                    if 'Heading 1' in style_name or 'Title' in style_name or 'Heading1' in style_name:
                        paragraphs_text.append(f"\n# {text}\n")
                    elif 'Heading 2' in style_name or 'Heading2' in style_name:
                        paragraphs_text.append(f"\n## {text}\n")
                    elif 'Heading 3' in style_name or 'Heading3' in style_name:
                        paragraphs_text.append(f"\n### {text}\n")
                    elif 'Heading 4' in style_name or 'Heading4' in style_name:
                        paragraphs_text.append(f"\n#### {text}\n")
                    else:
                        paragraphs_text.append(text)
                else:
                    paragraphs_text.append(text)
        
        # 处理未关闭的列表
        if in_list and current_list:
            lists_content.append(current_list)
        
        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # 合并单元格的内容会出现在第一个单元格中
                    cell_text = '\n'.join(extract_text_from_paragraph(p).strip() for p in cell.paragraphs).strip()
                    row_data.append(cell_text)
                if any(row_data):
                    table_data.append(row_data)
            if table_data:
                tables_content.append(table_data)
        
        # 合并所有内容
        full_content = '\n'.join(paragraphs_text)
        
        # 添加列表内容
        for lst in lists_content:
            for level, item in lst:
                indent = '  ' * level
                full_content += f"\n{indent}- {item}"
        
        # 添加表格内容（转换为markdown格式）
        for table in tables_content:
            if table:
                # 表头
                full_content += '\n\n| ' + ' | '.join(str(cell) for cell in table[0]) + ' |\n'
                # 分隔线
                full_content += '| ' + ' | '.join(['---'] * len(table[0])) + ' |\n'
                # 内容行
                for row in table[1:]:
                    full_content += '| ' + ' | '.join(str(cell) for cell in row) + ' |\n'
        
        return {
            "success": True,
            "error": "",
            "content": full_content.strip(),
            "structure": {
                "basic_info": [],
                "education": [],
                "work_experience": [],
                "projects": [],
                "skills": [],
                "certificates": [],
                "others": []
            },
            "stats": {
                "paragraphs": len(paragraphs_text),
                "tables": len(tables_content),
                "lists": len(lists_content)
            }
        }
    
    except Exception as e:
        return {"success": False, "error": f"解析文档时出错: {str(e)}", "content": ""}


def find_input_file(file_name):
    """在多个位置查找输入文件：技能根目录的 input/、当前运行目录的 input/、当前目录"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    skill_root_dir = os.path.dirname(script_dir)
    
    if os.path.exists(file_name):
        return file_name
    
    input_path = os.path.join(skill_root_dir, "input", file_name)
    if os.path.exists(input_path):
        return input_path
    
    input_path = os.path.join("input", file_name)
    if os.path.exists(input_path):
        return input_path
    
    return None


def main():
    if len(sys.argv) < 2:
        print("用法：python read_docx.py <input_file.docx> [output_file.txt]")
        print("      python read_docx.py <input_file.docx> --json [output_file.json]")
        print("      python read_docx.py <input_file.docx> --structure [output_file.json]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    resolved_input_file = find_input_file(input_file)
    if resolved_input_file is None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        skill_root_dir = os.path.dirname(script_dir)
        input_dir = os.path.join(skill_root_dir, "input")
        print(f"错误：找不到文件 '{input_file}'", file=sys.stderr)
        print(f"提示：请将文件放置在以下位置之一：", file=sys.stderr)
        print(f"  - 技能目录的 input/ 子目录：{input_dir}", file=sys.stderr)
        print(f"  - 当前运行目录的 input/ 子目录：{os.path.join(os.getcwd(), 'input')}", file=sys.stderr)
        print(f"  - 当前运行目录", file=sys.stderr)
        sys.exit(1)
    
    output_json = '--json' in sys.argv
    output_structure = '--structure' in sys.argv
    output_file = None
    
    for i, arg in enumerate(sys.argv):
        if arg in ('--json', '--structure'):
            if i + 1 < len(sys.argv) and not sys.argv[i + 1].startswith('--'):
                output_file = sys.argv[i + 1]
            break
        elif i == 2 and not arg.startswith('--'):
            output_file = arg
    
    result = parse_docx(resolved_input_file)
    
    if not result["success"]:
        print(f"错误：{result['error']}", file=sys.stderr)
        sys.exit(1)
    
    if output_json or output_structure:
        output_data = {
            "file": input_file,
            "content": result["content"],
            "structure": result["structure"],
            "stats": result["stats"]
        }
        json_str = json.dumps(output_data, ensure_ascii=False, indent=2)
        
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(json_str)
            print(f"JSON数据已保存到: {output_file}")
        else:
            print(json_str)
    else:
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(result["content"])
            print(f"文本内容已保存到: {output_file}")
        else:
            print(result["content"])


if __name__ == "__main__":
    main()