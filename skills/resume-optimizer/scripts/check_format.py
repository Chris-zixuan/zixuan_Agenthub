#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""检查Word文档run级别的行距"""

from docx import Document
import re

doc = Document('output/陈阳_简历.docx')

p = doc.paragraphs[8]
print(f"段落内容: {p.text[:50]}...")
print(f"段落中run的数量: {len(p.runs)}")

for k, run in enumerate(p.runs):
    print(f"\nRun {k}: {run.text[:30]}...")
    
    run_xml = run._r.xml
    if 'w:spacing' in run_xml:
        start = run_xml.find('<w:spacing')
        end = run_xml.find('/>', start) + 2
        print(f"  spacing: {run_xml[start:end]}")
        
        m = re.search(r'w:line="(\d+)"', run_xml[start:end])
        if m:
            line = int(m.group(1))
            line_mult = line / 240.0  # 240 twips = 1倍行距
            print(f"  行距值: {line} twips = {line_mult:.1f}倍")

print("\n" + "=" * 50)
print("检查二级标题（加粗）的行距:")

for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('工作经历'):
        print(f"\n二级标题: {p.text}")
        for k, run in enumerate(p.runs):
            run_xml = run._r.xml
            if 'w:spacing' in run_xml:
                start = run_xml.find('<w:spacing')
                end = run_xml.find('/>', start) + 2
                print(f"  spacing: {run_xml[start:end]}")
                
                m = re.search(r'w:line="(\d+)"', run_xml[start:end])
                if m:
                    line = int(m.group(1))
                    line_mult = line / 240.0
                    print(f"  行距值: {line} twips = {line_mult:.1f}倍")
        break
